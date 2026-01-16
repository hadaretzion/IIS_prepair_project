"""DOCX export service for final CV."""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re


def export_cv_to_docx(cv_text: str) -> bytes:
    """
    Convert CV text to DOCX format.
    
    Args:
        cv_text: CV text content
        
    Returns:
        DOCX file as bytes
    """
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    lines = cv_text.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
        
        # Detect section headers (all caps, short lines, or common section titles)
        is_header = False
        header_patterns = [
            r'^(SUMMARY|EXPERIENCE|EDUCATION|PROJECTS|SKILLS|CONTACT|PROFILE|OBJECTIVE)',
            r'^[A-Z][A-Z\s&]{2,30}$',  # All caps short lines
            r'^[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*:$'  # Title Case followed by colon
        ]
        
        for pattern in header_patterns:
            if re.match(pattern, line, re.IGNORECASE):
                is_header = True
                break
        
        if is_header:
            # Add as heading
            para = doc.add_heading(line, level=2)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            i += 1
        else:
            # Regular paragraph
            para = doc.add_paragraph(line)
            
            # Check if it's a bullet point
            if line.startswith('•') or line.startswith('-') or line.startswith('*'):
                para.style = 'List Bullet'
            elif re.match(r'^\d+[.)]\s', line):
                para.style = 'List Number'
            
            i += 1
    
    # Save to bytes
    import io
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
